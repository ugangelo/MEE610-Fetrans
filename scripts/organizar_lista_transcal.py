from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "MEE120_rev37_original.docx"
OUTPUT = ROOT / "MEE610_lista_exercicios_transcal.docx"


def element_text(element) -> str:
    return "".join(element.itertext()).strip()


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_exercise_number(paragraph, new_number: str) -> None:
    match = re.match(r"^(\d+\.\d+)(\s*[–-])", paragraph.text)
    if not match:
        return
    old_number = match.group(1)
    remaining = old_number
    inserted = False
    for run in paragraph.runs:
        if not remaining:
            return
        text = run.text
        if not text:
            continue
        take = min(len(text), len(remaining))
        if text[:take] != remaining[:take]:
            raise RuntimeError(f"Não foi possível renumerar: {paragraph.text[:80]}")
        replacement = new_number if not inserted else ""
        run.text = replacement + text[take:]
        inserted = True
        remaining = remaining[take:]
    if remaining:
        raise RuntimeError(f"Não foi possível renumerar: {paragraph.text[:80]}")


def organizar() -> None:
    doc = Document(SOURCE)

    # Identificação do documento, preservando o layout original da capa.
    cover = doc.tables[0]
    for cell in cover._cells:
        for paragraph in cell.paragraphs:
            if "MEE120" in paragraph.text:
                replace_paragraph_text(paragraph, paragraph.text.replace("MEE120", "MEE610"))

    footer = doc.sections[0].footer
    for paragraph in footer.paragraphs:
        if paragraph.text.strip():
            replace_paragraph_text(
                paragraph,
                paragraph.text.replace("MEE620", "MEE610").replace("V01", "V02"),
            )

    doc.core_properties.title = "MEE610 - Lista de exercícios de Transferência de Calor"
    doc.core_properties.subject = "Exercícios organizados conforme o cronograma de Transcal"

    body = doc._element.body
    children = list(body)

    headings = {}
    for index, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = element_text(child)
        for key in ("[2]", "[3]", "[4]", "[6]", "[8]", "[9]"):
            if text.startswith(key):
                headings[key] = index

    missing = [key for key in ("[2]", "[3]", "[4]", "[6]", "[8]", "[9]") if key not in headings]
    if missing:
        raise RuntimeError(f"Seções não encontradas: {missing}")

    starts = sorted((index, key) for key, index in headings.items())
    blocks = {}
    for position, (start, key) in enumerate(starts):
        end = starts[position + 1][0] if position + 1 < len(starts) else len(children)
        if children[end - 1].tag == qn("w:sectPr"):
            end -= 1
        blocks[key] = children[start:end]

    first_section = min(headings.values())
    sect_pr = next((child for child in children if child.tag == qn("w:sectPr")), None)

    for child in list(body)[first_section:]:
        body.remove(child)

    order = ("[2]", "[4]", "[3]", "[6]", "[8]", "[9]")
    for key in order:
        for child in blocks[key]:
            body.append(child)

    if sect_pr is not None:
        body.append(sect_pr)

    heading_names = {
        "[2]": "[1] – CONDUÇÃO UNIDIMENSIONAL PERMANENTE",
        "[4]": "[2] – PAREDES PLANAS, RESISTÊNCIAS TÉRMICAS E CONTATO",
        "[3]": "[3] – CONVECÇÃO E RADIAÇÃO",
        "[6]": "[4] – PAREDES CILÍNDRICAS E ESFÉRICAS / RAIO CRÍTICO",
        "[8]": "[5] – ALETAS",
        "[9]": "[6] – TROCADORES DE CALOR / DMLT",
    }
    current_section = None
    section_sequence = {"[2]": 1, "[4]": 2, "[3]": 3, "[6]": 4, "[8]": 5, "[9]": 6}
    exercise_counts = {number: 0 for number in section_sequence.values()}
    for paragraph in doc.paragraphs:
        for key, title in heading_names.items():
            if paragraph.text.startswith(key):
                current_section = section_sequence[key]
                replace_paragraph_text(paragraph, title)
                break
        else:
            if current_section is not None and re.match(r"^\d+\.\d+\s*[–-]", paragraph.text):
                exercise_counts[current_section] += 1
                replace_exercise_number(
                    paragraph,
                    f"{current_section}.{exercise_counts[current_section]}",
                )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    organizar()
