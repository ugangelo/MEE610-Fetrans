from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
PATH = ROOT / "MEE610_lista_exercicios_transcal.docx"


def main() -> None:
    doc = Document(PATH)
    removed = 0

    for paragraph in doc.paragraphs:
        # Mantém pageBreakBefore dos estilos de títulos; remove apenas quebras
        # explícitas inseridas em runs, que ficaram órfãs após a reordenação.
        for br in list(paragraph._element.iter(qn("w:br"))):
            if br.get(qn("w:type")) == "page":
                br.getparent().remove(br)
                removed += 1

    trailing = 0
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip() or paragraph._element.xpath(".//w:drawing|.//w:object|.//m:oMath"):
            break
        paragraph._element.getparent().remove(paragraph._element)
        trailing += 1

    doc.save(PATH)
    print(f"Quebras manuais removidas: {removed}")
    print(f"Parágrafos vazios finais removidos: {trailing}")


if __name__ == "__main__":
    main()
